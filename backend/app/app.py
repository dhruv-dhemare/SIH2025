

# perfectly working code for resume parsing (pdf/docx/images)

# import os
# import re
# import asyncio
# import pdfplumber
# from docx import Document
# from paddleocr import PaddleOCR
# from fastapi import FastAPI, File, UploadFile
# from fastapi.responses import JSONResponse
# import tempfile
# import httpx
# import json
# import spacy
# from dotenv import load_dotenv
# from PIL import Image, ImageEnhance, ImageFilter
# from fastapi import Form

# # ---------------------------
# # Load .env for API Key
# # ---------------------------
# load_dotenv()
# api_key = os.environ.get("GOOGLE_API_KEY")
# if not api_key:
#     raise ValueError("Set your GOOGLE_API_KEY in .env")

# # ---------------------------
# # App + Models Init
# # ---------------------------
# app = FastAPI()
# ocr = PaddleOCR(use_angle_cls=True, lang='en')  # robust for rotated text
# nlp = spacy.load("en_core_web_sm")

# # ---------------------------
# # Helpers
# # ---------------------------
# def get_file_extension(filename: str):
#     return filename.split(".")[-1].lower()

# def normalize_list(values):
#     seen, cleaned = set(), []
#     for v in values:
#         val = v.strip()
#         if not val or val.lower() in seen:
#             continue
#         seen.add(val.lower())
#         cleaned.append(val)
#     return cleaned

# # ---------------------------
# # Extract Text
# # ---------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     text = []
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#     return "\n".join(text)

# def extract_text_from_docx(file_path: str) -> str:
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# # ---------------------------
# # Image OCR Preprocessing
# # ---------------------------
# def _preprocess_image_for_ocr(path: str) -> str:
#     img = Image.open(path).convert("RGB")
#     target_w = 1400
#     if img.width < target_w:
#         ratio = target_w / img.width
#         img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
#     gray = img.convert("L")
#     gray = ImageEnhance.Contrast(gray).enhance(1.6)
#     gray = gray.filter(ImageFilter.SMOOTH)
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
#     tmp_path = tmp.name
#     tmp.close()
#     gray.save(tmp_path)
#     return tmp_path

# def extract_text_from_image(file_path: str) -> str:
#     if not os.path.exists(file_path):
#         return f"ERROR: File not found at {file_path}"

#     pre_path = None
#     results = None
#     try:
#         pre_path = _preprocess_image_for_ocr(file_path)
#         results = ocr.ocr(pre_path)
#     except:
#         results = None

#     if not results or len(results) == 0:
#         try:
#             results = ocr.ocr(file_path)
#         except:
#             results = None

#     lines = []
#     if results:
#         try:
#             for block in results:
#                 if isinstance(block, (list, tuple)):
#                     for item in block:
#                         candidate = item[1]
#                         if isinstance(candidate, (list, tuple)):
#                             text = candidate[0]
#                         else:
#                             text = candidate
#                         if text and isinstance(text, str) and text.strip():
#                             lines.append(text.strip())
#                 else:
#                     s = str(block)
#                     if s.strip():
#                         lines.append(s.strip())
#         except:
#             lines = [str(results)]

#     if pre_path and os.path.exists(pre_path):
#         os.remove(pre_path)

#     if not lines:
#         return "ERROR: No text detected in image (PaddleOCR returned empty)."
#     return "\n".join(lines)

# # ---------------------------
# # Regex + spaCy Preprocessing
# # ---------------------------
# def extract_quick_fields(text: str):
#     emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
#     phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))
#     doc = nlp(text)
#     names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
#     orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
#     locations = normalize_list([ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]])
#     edu_matches = re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I)
#     education_list = normalize_list(edu_matches)
#     skills_matches = re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I)
#     skills_list = normalize_list(skills_matches)
#     project_matches = re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text)
#     projects_list = normalize_list(project_matches)
#     experience_matches = re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text)
#     experience_list = normalize_list(["{} {} {}".format(*exp).strip() for exp in experience_matches if any(exp)])
#     achievement_matches = re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text)
#     achievements_list = normalize_list(achievement_matches)
#     extracurricular_matches = re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text)
#     extracurricular_list = normalize_list(extracurricular_matches)

#     return {
#         "emails": emails,
#         "phones": phones,
#         "names": names,
#         "orgs": orgs,
#         "locations": locations,
#         "education": education_list,
#         "skills": skills_list,
#         "projects": projects_list,
#         "experience": experience_list,
#         "achievements": achievements_list,
#         "extracurricular": extracurricular_list
#     }

# # ---------------------------
# # Gemini Cleaning (optional, not parsing)
# # ---------------------------
# async def clean_with_gemini(raw_json: dict):
#     url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
#     prompt = f"""
# You are a resume cleaning assistant. Input is already parsed JSON.
# Remove duplicates, irrelevant values, normalize names/degrees/skills.
# Return only valid JSON, no markdown or ```json fences.
# JSON:
# {json.dumps(raw_json, indent=2)}
# """
#     payload = {"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature":0, "maxOutputTokens":2000}}
#     async with httpx.AsyncClient(timeout=60) as client:
#         response = await client.post(url, json=payload)
#         response.raise_for_status()
#         result = response.json()
#     candidates = result.get("candidates", [])
#     if not candidates:
#         return {"error": "No candidates returned from Gemini"}
#     parts = candidates[0].get("content", {}).get("parts", [])
#     content = parts[0].get("text", "") if parts else ""
#     content = content.strip()
#     if content.startswith("```"):
#         content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()
#     try:
#         return json.loads(content)
#     except:
#         return {"error": "Failed to parse cleaned JSON", "raw": content}

# # ---------------------------
# # API Endpoints
# # ---------------------------
# @app.post("/ocr-debug")
# async def ocr_debug(file: UploadFile = File(...)):
#     ext = get_file_extension(file.filename)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#         tmp.write(await file.read())
#         tmp_path = tmp.name
#     try:
#         raw = ocr.ocr(tmp_path)
#         os.remove(tmp_path)
#         return {"status": "ok", "ocr_raw": str(raw)}
#     except Exception as e:
#         try: os.remove(tmp_path)
#         except: pass
#         return {"status": "error", "message": str(e)}

# @app.post("/parse-resume")
# async def parse_resume_file(file: UploadFile = File(...)):
#     try:
#         ext = get_file_extension(file.filename)
#         with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#             tmp.write(await file.read())
#             tmp_path = tmp.name

#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png", "jpg", "jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             os.remove(tmp_path)
#             return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

#         os.remove(tmp_path)
#         extracted_data = extract_quick_fields(text)

#         raw_json = {
#             "personal_info": {
#                 "name": extracted_data["names"][0] if extracted_data["names"] else None,
#                 "emails": extracted_data["emails"],
#                 "phones": extracted_data["phones"],
#                 "locations": extracted_data["locations"],
#                 "organizations": extracted_data["orgs"]
#             },
#             "education": extracted_data["education"],
#             "experience": extracted_data["experience"],
#             "skills": extracted_data["skills"],
#             "projects": extracted_data["projects"],
#             "achievements": extracted_data["achievements"],
#             "extracurricular": extracted_data["extracurricular"]
#         }

#         structured_data = await clean_with_gemini(raw_json)
#         return {"status": "ok", "resume_parsed": structured_data, "scanned": ext in ["png", "jpg", "jpeg"]}
#     except Exception as e:
#         return JSONResponse(status_code=500, content={"status": "error", "message": str(e)})



# @app.post("/signup/")
# async def signup_user(
#     name: str = Form(...),
#     email: str = Form(...),
#     password: str = Form(...),
#     role: str = Form(...),
#     file: UploadFile = None
# ):
#     user_data = {
#         "name": name,
#         "email": email,
#         "password": password,  # hash it in real app!
#         "role": role
#     }

#     parsed_resume = None
#     if file:
#         # Save file temporarily
#         ext = get_file_extension(file.filename)
#         with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#             tmp.write(await file.read())
#             tmp_path = tmp.name

#         # Parse resume using your pipeline
#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png", "jpg", "jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             os.remove(tmp_path)
#             return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

#         os.remove(tmp_path)
#         extracted_data = extract_quick_fields(text)

#         raw_json = {
#             "personal_info": {
#                 "name": extracted_data["names"][0] if extracted_data["names"] else None,
#                 "emails": extracted_data["emails"],
#                 "phones": extracted_data["phones"],
#                 "locations": extracted_data["locations"],
#                 "organizations": extracted_data["orgs"]
#             },
#             "education": extracted_data["education"],
#             "experience": extracted_data["experience"],
#             "skills": extracted_data["skills"],
#             "projects": extracted_data["projects"],
#             "achievements": extracted_data["achievements"],
#             "extracurricular": extracted_data["extracurricular"]
#         }

#         parsed_resume = await clean_with_gemini(raw_json)

#     # Save user and parsed resume to DB
#     # Example: you can append to a JSON file or DB table
#     # For now, we just return it
#     user_data["resume_parsed"] = parsed_resume

#     return {"status": "ok", "message": "Signup successful", "user": user_data}


#Integrated with sign up page Auto filling registration form works

import os
import re
import asyncio
import pdfplumber
from docx import Document
from paddleocr import PaddleOCR
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
import tempfile
import httpx
import json
import spacy
from dotenv import load_dotenv
from PIL import Image, ImageEnhance, ImageFilter
from fastapi import Form

from typing import Optional, List
import json


# At the top of app.py
from fastapi.middleware.cors import CORSMiddleware
app = FastAPI()

origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,  # allow these origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------
# Load .env for API Key
# ---------------------------
load_dotenv()
api_key = os.environ.get("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("Set your GOOGLE_API_KEY in .env")


# ---------------------------
# App + Models Init
# ---------------------------

ocr = PaddleOCR(use_angle_cls=True, lang='en')  # robust for rotated text
nlp = spacy.load("en_core_web_sm")

# ---------------------------
# Helpers
# ---------------------------
def get_file_extension(filename: str):
    return filename.split(".")[-1].lower()

# async def save_upload_file(file: UploadFile, folder: str, prefix: str = "") -> str:
#     filename = f"{prefix}_{file.filename}"
#     file_path = os.path.join(folder, filename)
#     with open(file_path, "wb") as f:
#         f.write(await file.read())
#     return file_path

def normalize_list(values):
    seen, cleaned = set(), []
    for v in values:
        val = v.strip()
        if not val or val.lower() in seen:
            continue
        seen.add(val.lower())
        cleaned.append(val)
    return cleaned

# ---------------------------
# Extract Text
# ---------------------------
def extract_text_from_pdf(file_path: str) -> str:
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# ---------------------------
# Image OCR Preprocessing
# ---------------------------
def _preprocess_image_for_ocr(path: str) -> str:
    img = Image.open(path).convert("RGB")
    target_w = 1400
    if img.width < target_w:
        ratio = target_w / img.width
        img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.6)
    gray = gray.filter(ImageFilter.SMOOTH)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = tmp.name
    tmp.close()
    gray.save(tmp_path)
    return tmp_path

def extract_text_from_image(file_path: str) -> str:
    if not os.path.exists(file_path):
        return f"ERROR: File not found at {file_path}"

    pre_path = None
    results = None
    try:
        pre_path = _preprocess_image_for_ocr(file_path)
        results = ocr.ocr(pre_path)
    except:
        results = None

    if not results or len(results) == 0:
        try:
            results = ocr.ocr(file_path)
        except:
            results = None

    lines = []
    if results:
        try:
            for block in results:
                if isinstance(block, (list, tuple)):
                    for item in block:
                        candidate = item[1]
                        if isinstance(candidate, (list, tuple)):
                            text = candidate[0]
                        else:
                            text = candidate
                        if text and isinstance(text, str) and text.strip():
                            lines.append(text.strip())
                else:
                    s = str(block)
                    if s.strip():
                        lines.append(s.strip())
        except:
            lines = [str(results)]

    if pre_path and os.path.exists(pre_path):
        os.remove(pre_path)

    if not lines:
        return "ERROR: No text detected in image (PaddleOCR returned empty)."
    return "\n".join(lines)

# ---------------------------
# Regex + spaCy Preprocessing
# ---------------------------
def extract_quick_fields(text: str):
    emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
    phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))
    doc = nlp(text)
    names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
    orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
    locations = normalize_list([ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]])
    edu_matches = re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I)
    education_list = normalize_list(edu_matches)
    skills_matches = re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I)
    skills_list = normalize_list(skills_matches)
    project_matches = re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text)
    projects_list = normalize_list(project_matches)
    experience_matches = re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text)
    experience_list = normalize_list(["{} {} {}".format(*exp).strip() for exp in experience_matches if any(exp)])
    achievement_matches = re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text)
    achievements_list = normalize_list(achievement_matches)
    extracurricular_matches = re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text)
    extracurricular_list = normalize_list(extracurricular_matches)

    return {
        "emails": emails,
        "phones": phones,
        "names": names,
        "orgs": orgs,
        "locations": locations,
        "education": education_list,
        "skills": skills_list,
        "projects": projects_list,
        "experience": experience_list,
        "achievements": achievements_list,
        "extracurricular": extracurricular_list
    }

# ---------------------------
# Gemini Cleaning (optional, not parsing)
# ---------------------------
async def clean_with_gemini(raw_json: dict):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    prompt = f"""
You are a resume cleaning assistant. Input is already parsed JSON.
Remove duplicates, irrelevant values, normalize names/degrees/skills.
Return only valid JSON, no markdown or ```json fences.
JSON:
{json.dumps(raw_json, indent=2)}
"""
    payload = {"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature":0, "maxOutputTokens":2000}}
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(url, json=payload)
        response.raise_for_status()
        result = response.json()
    candidates = result.get("candidates", [])
    if not candidates:
        return {"error": "No candidates returned from Gemini"}
    parts = candidates[0].get("content", {}).get("parts", [])
    content = parts[0].get("text", "") if parts else ""
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()
    try:
        return json.loads(content)
    except:
        return {"error": "Failed to parse cleaned JSON", "raw": content}

# ---------------------------
# API Endpoints
# ---------------------------
@app.post("/ocr-debug")
async def ocr_debug(file: UploadFile = File(...)):
    ext = get_file_extension(file.filename)
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        raw = ocr.ocr(tmp_path)
        os.remove(tmp_path)
        return {"status": "ok", "ocr_raw": str(raw)}
    except Exception as e:
        try: os.remove(tmp_path)
        except: pass
        return {"status": "error", "message": str(e)}

@app.post("/parse-resume")
async def parse_resume_file(file: UploadFile = File(...)):
    try:
        ext = get_file_extension(file.filename)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        if ext == "pdf":
            text = extract_text_from_pdf(tmp_path)
        elif ext == "docx":
            text = extract_text_from_docx(tmp_path)
        elif ext in ["png", "jpg", "jpeg"]:
            text = extract_text_from_image(tmp_path)
        else:
            os.remove(tmp_path)
            return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

        os.remove(tmp_path)
        extracted_data = extract_quick_fields(text)

        raw_json = {
            "personal_info": {
                "name": extracted_data["names"][0] if extracted_data["names"] else None,
                "emails": extracted_data["emails"],
                "phones": extracted_data["phones"],
                "locations": extracted_data["locations"],
                "organizations": extracted_data["orgs"]
            },
            "education": extracted_data["education"],
            "experience": extracted_data["experience"],
            "skills": extracted_data["skills"],
            "projects": extracted_data["projects"],
            "achievements": extracted_data["achievements"],
            "extracurricular": extracted_data["extracurricular"]
        }

        structured_data = await clean_with_gemini(raw_json)
        return {"status": "ok", "resume_parsed": structured_data, "scanned": ext in ["png", "jpg", "jpeg"]}
    except Exception as e:
        return JSONResponse(status_code=500, content={"status": "error", "message": str(e)})


# Add this to your existing FastAPI code

# Update your existing signup endpoint to handle the parsed resume data

# Add these endpoints to your existing app.py file (after your existing code)



# ============= SIGNUP ENDPOINTS =============

@app.post("/api/student/signup")
async def signup_student(
    # Basic required fields
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    
    # Optional fields
    phn: Optional[str] = Form(None),
    headline: Optional[str] = Form(None),
    about: Optional[str] = Form(None),
    department: Optional[str] = Form(None),
    yearOfStudy: Optional[str] = Form(None),
    
    # Address fields
    address: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    state: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    dob: Optional[str] = Form(None),
    gender: Optional[str] = Form(None),
    
    # Array fields - these come as strings, need to be split
    experience: Optional[str] = Form(None),
    certification: Optional[str] = Form(None),
    skills: Optional[str] = Form(None),
    
    # List fields - these come as JSON strings
    education: Optional[str] = Form(None),  # JSON string of array
    urls: Optional[str] = Form(None),       # JSON string of array
    locations: Optional[str] = Form(None),  # JSON string of array
    
    # Files
    resume: UploadFile = File(None),
    profilePhoto: UploadFile = File(None),
    
    # Resume parsing flag
    resumeParsed: Optional[bool] = Form(False)
):
    try:
        # Process array fields
        experience_list = experience.split(',') if experience else []
        experience_list = [exp.strip() for exp in experience_list if exp.strip()]
        
        certification_list = certification.split(',') if certification else []
        certification_list = [cert.strip() for cert in certification_list if cert.strip()]
        
        skills_list = skills.split(',') if skills else []
        skills_list = [skill.strip() for skill in skills_list if skill.strip()]
        
        # Process JSON fields
        try:
            education_list = json.loads(education) if education else []
        except:
            education_list = [education] if education else []
            
        try:
            urls_list = json.loads(urls) if urls else []
        except:
            urls_list = [urls] if urls else []
            
        try:
            locations_list = json.loads(locations) if locations else []
        except:
            locations_list = []
        
        # Create user data structure
        user_data = {
            "name": name,
            "email": email,
            "password": password,  # In production: hash this!
            "role": role,
            "phn": phn,
            "headline": headline,
            "about": about,
            "department": department,
            "yearOfStudy": yearOfStudy,
            "address": address,
            "city": city,
            "state": state,
            "country": country,
            "dob": dob,
            "gender": gender,
            "experience": experience_list,
            "certification": certification_list,
            "skills": skills_list,
            "education": education_list,
            "urls": urls_list,
            "locations": locations_list,
            "resumeParsed": resumeParsed,
            "files": {}
        }
        
        # Handle file uploads
        if resume:
            # Save resume file
            resume_filename = f"resume_{email}_{resume.filename}"
            # TODO: Implement your file saving logic here
            user_data["files"]["resume"] = resume_filename
            
        if profilePhoto:
            # Save profile photo
            photo_filename = f"photo_{email}_{profilePhoto.filename}"
            # TODO: Implement your file saving logic here
            user_data["files"]["profilePhoto"] = photo_filename
        
        # TODO: Save to your database here
        # Example: save_student_to_db(user_data)
        
        return {
            "status": "success",
            "message": "Student registered successfully!",
            "user_id": email,  # or generate proper ID
            "data": user_data
        }
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"Registration failed: {str(e)}"}
        )

@app.post("/api/alumni/signup")
async def signup_alumni(
    # Same parameters as student
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    phn: Optional[str] = Form(None),
    headline: Optional[str] = Form(None),
    about: Optional[str] = Form(None),
    department: Optional[str] = Form(None),
    yearOfStudy: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    state: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    dob: Optional[str] = Form(None),
    gender: Optional[str] = Form(None),
    experience: Optional[str] = Form(None),
    certification: Optional[str] = Form(None),
    skills: Optional[str] = Form(None),
    education: Optional[str] = Form(None),
    urls: Optional[str] = Form(None),
    locations: Optional[str] = Form(None),
    resume: UploadFile = File(None),
    profilePhoto: UploadFile = File(None),
    resumeParsed: Optional[bool] = Form(False)
):
    # Same logic as student signup
    return await signup_student(
        name, email, password, role, phn, headline, about, department, 
        yearOfStudy, address, city, state, country, dob, gender,
        experience, certification, skills, education, urls, locations,
        resume, profilePhoto, resumeParsed
    )

@app.post("/api/recruiter/signup")
async def signup_recruiter(
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    phn: Optional[str] = Form(None),
    about: Optional[str] = Form(None),
    company: Optional[str] = Form(None),
    position: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    state: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    dob: Optional[str] = Form(None),
    gender: Optional[str] = Form(None),
    urls: Optional[str] = Form(None),
    locations: Optional[str] = Form(None),
    profilePhoto: UploadFile = File(None),
):
    try:
        # Process JSON fields
        try:
            urls_list = json.loads(urls) if urls else []
        except:
            urls_list = [urls] if urls else []
            
        try:
            locations_list = json.loads(locations) if locations else []
        except:
            locations_list = []
        
        user_data = {
            "name": name,
            "email": email,
            "password": password,  # In production: hash this!
            "role": role,
            "phn": phn,
            "about": about,
            "company": company,
            "position": position,
            "address": address,
            "city": city,
            "state": state,
            "country": country,
            "dob": dob,
            "gender": gender,
            "urls": urls_list,
            "locations": locations_list,
            "files": {}
        }
        
        if profilePhoto:
            photo_filename = f"photo_{email}_{profilePhoto.filename}"
            user_data["files"]["profilePhoto"] = photo_filename
        
        # TODO: Save to your database here
        
        return {
            "status": "success",
            "message": "Recruiter registered successfully!",
            "user_id": email,
            "data": user_data
        }
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"Registration failed: {str(e)}"}
        )

@app.post("/api/teacher/signup")
async def signup_teacher(
    # Same as alumni/student for faculty
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    phn: Optional[str] = Form(None),
    headline: Optional[str] = Form(None),
    about: Optional[str] = Form(None),
    department: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    state: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    dob: Optional[str] = Form(None),
    gender: Optional[str] = Form(None),
    experience: Optional[str] = Form(None),
    certification: Optional[str] = Form(None),
    skills: Optional[str] = Form(None),
    education: Optional[str] = Form(None),
    urls: Optional[str] = Form(None),
    locations: Optional[str] = Form(None),
    resume: UploadFile = File(None),
    profilePhoto: UploadFile = File(None),
    resumeParsed: Optional[bool] = Form(False)
):
    return await signup_student(
        name, email, password, role, phn, headline, about, department, 
        None, address, city, state, country, dob, gender,  # yearOfStudy = None for teachers
        experience, certification, skills, education, urls, locations,
        resume, profilePhoto, resumeParsed
    )

@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "message": "Backend running on port 8000",
        "endpoints": ["/parse-resume", "/health"]
    }
