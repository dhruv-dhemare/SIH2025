

# perfectly working code for resume parsing (pdf/docx/images)

# import os
# import re
# import asyncio
# import pdfplumber
# from docx import Document
# from paddleocr import PaddleOCR
# from fastapi import FastAPI, File, UploadFile
# from fastapi.responses import JSONResponse
# import tempfile
# import httpx
# import json
# import spacy
# from dotenv import load_dotenv
# from PIL import Image, ImageEnhance, ImageFilter
# from fastapi import Form

# # ---------------------------
# # Load .env for API Key
# # ---------------------------
# load_dotenv()
# api_key = os.environ.get("GOOGLE_API_KEY")
# if not api_key:
#     raise ValueError("Set your GOOGLE_API_KEY in .env")

# # ---------------------------
# # App + Models Init
# # ---------------------------
# app = FastAPI()
# ocr = PaddleOCR(use_angle_cls=True, lang='en')  # robust for rotated text
# nlp = spacy.load("en_core_web_sm")

# # ---------------------------
# # Helpers
# # ---------------------------
# def get_file_extension(filename: str):
#     return filename.split(".")[-1].lower()

# def normalize_list(values):
#     seen, cleaned = set(), []
#     for v in values:
#         val = v.strip()
#         if not val or val.lower() in seen:
#             continue
#         seen.add(val.lower())
#         cleaned.append(val)
#     return cleaned

# # ---------------------------
# # Extract Text
# # ---------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     text = []
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#     return "\n".join(text)

# def extract_text_from_docx(file_path: str) -> str:
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# # ---------------------------
# # Image OCR Preprocessing
# # ---------------------------
# def _preprocess_image_for_ocr(path: str) -> str:
#     img = Image.open(path).convert("RGB")
#     target_w = 1400
#     if img.width < target_w:
#         ratio = target_w / img.width
#         img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
#     gray = img.convert("L")
#     gray = ImageEnhance.Contrast(gray).enhance(1.6)
#     gray = gray.filter(ImageFilter.SMOOTH)
#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
#     tmp_path = tmp.name
#     tmp.close()
#     gray.save(tmp_path)
#     return tmp_path

# def extract_text_from_image(file_path: str) -> str:
#     if not os.path.exists(file_path):
#         return f"ERROR: File not found at {file_path}"

#     pre_path = None
#     results = None
#     try:
#         pre_path = _preprocess_image_for_ocr(file_path)
#         results = ocr.ocr(pre_path)
#     except:
#         results = None

#     if not results or len(results) == 0:
#         try:
#             results = ocr.ocr(file_path)
#         except:
#             results = None

#     lines = []
#     if results:
#         try:
#             for block in results:
#                 if isinstance(block, (list, tuple)):
#                     for item in block:
#                         candidate = item[1]
#                         if isinstance(candidate, (list, tuple)):
#                             text = candidate[0]
#                         else:
#                             text = candidate
#                         if text and isinstance(text, str) and text.strip():
#                             lines.append(text.strip())
#                 else:
#                     s = str(block)
#                     if s.strip():
#                         lines.append(s.strip())
#         except:
#             lines = [str(results)]

#     if pre_path and os.path.exists(pre_path):
#         os.remove(pre_path)

#     if not lines:
#         return "ERROR: No text detected in image (PaddleOCR returned empty)."
#     return "\n".join(lines)

# # ---------------------------
# # Regex + spaCy Preprocessing
# # ---------------------------
# def extract_quick_fields(text: str):
#     emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
#     phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))
#     doc = nlp(text)
#     names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
#     orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
#     locations = normalize_list([ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]])
#     edu_matches = re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I)
#     education_list = normalize_list(edu_matches)
#     skills_matches = re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I)
#     skills_list = normalize_list(skills_matches)
#     project_matches = re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text)
#     projects_list = normalize_list(project_matches)
#     experience_matches = re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text)
#     experience_list = normalize_list(["{} {} {}".format(*exp).strip() for exp in experience_matches if any(exp)])
#     achievement_matches = re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text)
#     achievements_list = normalize_list(achievement_matches)
#     extracurricular_matches = re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text)
#     extracurricular_list = normalize_list(extracurricular_matches)

#     return {
#         "emails": emails,
#         "phones": phones,
#         "names": names,
#         "orgs": orgs,
#         "locations": locations,
#         "education": education_list,
#         "skills": skills_list,
#         "projects": projects_list,
#         "experience": experience_list,
#         "achievements": achievements_list,
#         "extracurricular": extracurricular_list
#     }

# # ---------------------------
# # Gemini Cleaning (optional, not parsing)
# # ---------------------------
# async def clean_with_gemini(raw_json: dict):
#     url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
#     prompt = f"""
# You are a resume cleaning assistant. Input is already parsed JSON.
# Remove duplicates, irrelevant values, normalize names/degrees/skills.
# Return only valid JSON, no markdown or ```json fences.
# JSON:
# {json.dumps(raw_json, indent=2)}
# """
#     payload = {"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature":0, "maxOutputTokens":2000}}
#     async with httpx.AsyncClient(timeout=60) as client:
#         response = await client.post(url, json=payload)
#         response.raise_for_status()
#         result = response.json()
#     candidates = result.get("candidates", [])
#     if not candidates:
#         return {"error": "No candidates returned from Gemini"}
#     parts = candidates[0].get("content", {}).get("parts", [])
#     content = parts[0].get("text", "") if parts else ""
#     content = content.strip()
#     if content.startswith("```"):
#         content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()
#     try:
#         return json.loads(content)
#     except:
#         return {"error": "Failed to parse cleaned JSON", "raw": content}

# # ---------------------------
# # API Endpoints
# # ---------------------------
# @app.post("/ocr-debug")
# async def ocr_debug(file: UploadFile = File(...)):
#     ext = get_file_extension(file.filename)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#         tmp.write(await file.read())
#         tmp_path = tmp.name
#     try:
#         raw = ocr.ocr(tmp_path)
#         os.remove(tmp_path)
#         return {"status": "ok", "ocr_raw": str(raw)}
#     except Exception as e:
#         try: os.remove(tmp_path)
#         except: pass
#         return {"status": "error", "message": str(e)}

# @app.post("/parse-resume")
# async def parse_resume_file(file: UploadFile = File(...)):
#     try:
#         ext = get_file_extension(file.filename)
#         with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#             tmp.write(await file.read())
#             tmp_path = tmp.name

#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png", "jpg", "jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             os.remove(tmp_path)
#             return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

#         os.remove(tmp_path)
#         extracted_data = extract_quick_fields(text)

#         raw_json = {
#             "personal_info": {
#                 "name": extracted_data["names"][0] if extracted_data["names"] else None,
#                 "emails": extracted_data["emails"],
#                 "phones": extracted_data["phones"],
#                 "locations": extracted_data["locations"],
#                 "organizations": extracted_data["orgs"]
#             },
#             "education": extracted_data["education"],
#             "experience": extracted_data["experience"],
#             "skills": extracted_data["skills"],
#             "projects": extracted_data["projects"],
#             "achievements": extracted_data["achievements"],
#             "extracurricular": extracted_data["extracurricular"]
#         }

#         structured_data = await clean_with_gemini(raw_json)
#         return {"status": "ok", "resume_parsed": structured_data, "scanned": ext in ["png", "jpg", "jpeg"]}
#     except Exception as e:
#         return JSONResponse(status_code=500, content={"status": "error", "message": str(e)})



# @app.post("/signup/")
# async def signup_user(
#     name: str = Form(...),
#     email: str = Form(...),
#     password: str = Form(...),
#     role: str = Form(...),
#     file: UploadFile = None
# ):
#     user_data = {
#         "name": name,
#         "email": email,
#         "password": password,  # hash it in real app!
#         "role": role
#     }

#     parsed_resume = None
#     if file:
#         # Save file temporarily
#         ext = get_file_extension(file.filename)
#         with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#             tmp.write(await file.read())
#             tmp_path = tmp.name

#         # Parse resume using your pipeline
#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png", "jpg", "jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             os.remove(tmp_path)
#             return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

#         os.remove(tmp_path)
#         extracted_data = extract_quick_fields(text)

#         raw_json = {
#             "personal_info": {
#                 "name": extracted_data["names"][0] if extracted_data["names"] else None,
#                 "emails": extracted_data["emails"],
#                 "phones": extracted_data["phones"],
#                 "locations": extracted_data["locations"],
#                 "organizations": extracted_data["orgs"]
#             },
#             "education": extracted_data["education"],
#             "experience": extracted_data["experience"],
#             "skills": extracted_data["skills"],
#             "projects": extracted_data["projects"],
#             "achievements": extracted_data["achievements"],
#             "extracurricular": extracted_data["extracurricular"]
#         }

#         parsed_resume = await clean_with_gemini(raw_json)

#     # Save user and parsed resume to DB
#     # Example: you can append to a JSON file or DB table
#     # For now, we just return it
#     user_data["resume_parsed"] = parsed_resume

#     return {"status": "ok", "message": "Signup successful", "user": user_data}
# backend/main.py
import os
import re
import tempfile
import json
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import pdfplumber
from paddleocr import PaddleOCR
from PIL import Image, ImageEnhance, ImageFilter
import spacy
import httpx
from pathlib import Path
from dotenv import load_dotenv

# ---------------------------
# Load .env
# ---------------------------
env_path = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=env_path)

api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("Set your GOOGLE_API_KEY in backend/.env")


# ---------------------------
# Directories
# ---------------------------
UPLOAD_DIR = "uploads"
PARSED_DIR = "parsed_resume"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PARSED_DIR, exist_ok=True)

# ---------------------------
# App Init
# ---------------------------
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# NLP & OCR
ocr = PaddleOCR(use_angle_cls=True, lang='en')
nlp = spacy.load("en_core_web_sm")

# ---------------------------
# Helpers
# ---------------------------
def get_file_extension(filename: str):
    return filename.split(".")[-1].lower()

def normalize_list(values):
    seen, cleaned = set(), []
    for v in values:
        # If v is a tuple (from regex with multiple groups), join parts into a string
        if isinstance(v, tuple):
            val = " ".join([str(part).strip() for part in v if part])
        else:
            val = str(v).strip()

        if not val or val.lower() in seen:
            continue
        seen.add(val.lower())
        cleaned.append(val)
    return cleaned


# ---------------------------
# Text extraction
# ---------------------------
def extract_text_from_pdf(path: str) -> str:
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: text.append(t)
    return "\n".join(text)

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def preprocess_image(path: str) -> str:
    img = Image.open(path).convert("RGB")
    target_w = 1400
    if img.width < target_w:
        ratio = target_w / img.width
        img = img.resize((int(img.width*ratio), int(img.height*ratio)), Image.LANCZOS)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.6)
    gray = gray.filter(ImageFilter.SMOOTH)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = tmp.name
    tmp.close()
    gray.save(tmp_path)
    return tmp_path

def extract_text_from_image(path: str) -> str:
    pre_path = None
    try:
        pre_path = preprocess_image(path)
        result = ocr.ocr(pre_path)
    except:
        result = ocr.ocr(path)
    finally:
        if pre_path and os.path.exists(pre_path):
            os.remove(pre_path)

    lines = []
    if result:
        for block in result:
            if isinstance(block, (list, tuple)):
                for item in block:
                    text = item[1][0] if isinstance(item[1], (list, tuple)) else item[1]
                    if text.strip(): lines.append(text.strip())
    return "\n".join(lines)

# ---------------------------
# Resume field extraction
# ---------------------------
def extract_quick_fields(text: str):
    emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
    phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))
    doc = nlp(text)
    names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
    orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
    edu_matches = normalize_list(re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I))
    skills_matches = normalize_list(re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I))
    project_matches = normalize_list(re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text))
    experience_matches = normalize_list(re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text))
    achievements_list = normalize_list(re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text))
    extracurricular_list = normalize_list(re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text))

    return {
        "emails": emails, "phones": phones, "names": names, "orgs": orgs, "education": edu_matches, "skills": skills_matches,
        "projects": project_matches, "experience": ["{} {} {}".format(*(list(exp) + [""] * (3 - len(exp)))).strip() for exp in experience_matches if any(exp)],
        "achievements": achievements_list, "extracurricular": extracurricular_list
    }

# ---------------------------
# Gemini cleaning
# ---------------------------
async def clean_with_gemini(raw_json: dict):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    prompt = f"""
Clean resume JSON. Remove duplicates, irrelevant values, normalize names/degrees/skills.
Return only valid JSON, no markdown.
Input JSON:
{json.dumps(raw_json, indent=2)}
"""
    payload = {"contents":[{"parts":[{"text": prompt}]}],"generationConfig":{"temperature":0,"maxOutputTokens":2000}}
    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(url, json=payload)
        resp.raise_for_status()
        data = resp.json()
    candidates = data.get("candidates", [])
    if not candidates: return {"error": "No candidates returned"}
    parts = candidates[0].get("content", {}).get("parts", [])
    content = parts[0].get("text","") if parts else ""
    content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()
    try:
        return json.loads(content)
    except:
        return {"error": "Failed to parse cleaned JSON", "raw": content}

# ---------------------------
# In-memory DB
# ---------------------------
USERS_DB = {}

# ---------------------------
# Routes
# ---------------------------
@app.post("/signup")
async def signup_user(
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    file: UploadFile = None
):
    user_data = {"name": name, "email": email, "password": password, "role": role}
    parsed_resume = None

    if file:
        ext = get_file_extension(file.filename)
        save_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(save_path, "wb") as f:
            f.write(await file.read())

        if ext == "pdf":
            text = extract_text_from_pdf(save_path)
        elif ext == "docx":
            text = extract_text_from_docx(save_path)
        elif ext in ["png","jpg","jpeg"]:
            text = extract_text_from_image(save_path)
        else:
            return JSONResponse(status_code=400, content={"status":"error", "message": f"Unsupported file type: {ext}"})

        extracted_data = extract_quick_fields(text)
        raw_json = {
            "personal_info": {
                "name": extracted_data["names"][0] if extracted_data.get("names") else None,
                "emails": extracted_data.get("emails", []),
                "phones": extracted_data.get("phones", []),
                "organizations": extracted_data.get("orgs", [])
            },
            "education": extracted_data.get("education", []),
            "experience": extracted_data.get("experience", []),
            "skills": extracted_data.get("skills", []),
            "projects": extracted_data.get("projects", []),
            "achievements": extracted_data.get("achievements", []),
            "extracurricular": extracted_data.get("extracurricular", [])
        }

        parsed_resume = await clean_with_gemini(raw_json)

        parsed_path = os.path.join(PARSED_DIR, f"{file.filename}.json")
        with open(parsed_path, "w", encoding="utf-8") as f:
            json.dump(parsed_resume, f, ensure_ascii=False, indent=2)

    user_data["resume_parsed"] = parsed_resume
    USERS_DB[email] = user_data
    return {"status":"ok", "message":"Signup successful", "user": user_data}


@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    ext = get_file_extension(file.filename)
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        if ext == "pdf":
            text = extract_text_from_pdf(tmp_path)
        elif ext == "docx":
            text = extract_text_from_docx(tmp_path)
        elif ext in ["png","jpg","jpeg"]:
            text = extract_text_from_image(tmp_path)
        else:
            return JSONResponse(status_code=400, content={"status":"error", "message": f"Unsupported file type: {ext}"})
    finally:
        os.remove(tmp_path)

    extracted = extract_quick_fields(text)
    raw_json = {
        "personal_info": {
            "name": extracted["names"][0] if extracted.get("names") else None,
            "emails": extracted.get("emails", []),
            "phones": extracted.get("phones", []),
            "organizations": extracted.get("orgs", [])
        },
        "education": extracted.get("education", []),
        "experience": extracted.get("experience", []),
        "skills": extracted.get("skills", []),
        "projects": extracted.get("projects", []),
        "achievements": extracted.get("achievements", []),
        "extracurricular": extracted.get("extracurricular", [])
    }

    structured = await clean_with_gemini(raw_json)
    parsed_path = os.path.join(PARSED_DIR, f"{file.filename}.json")
    with open(parsed_path, "w", encoding="utf-8") as f:
        json.dump(structured, f, ensure_ascii=False, indent=2)

    return {"status":"ok", "resume_parsed": structured, "scanned": ext in ["png","jpg","jpeg"]}

# @app.post("/parse-resume")
# async def parse_resume(file: UploadFile = File(...)):
#     ext = get_file_extension(file.filename)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#         tmp.write(await file.read())
#         tmp_path = tmp.name

#     try:
#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png","jpg","jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             return JSONResponse(status_code=400, content={"status":"error", "message": f"Unsupported file type: {ext}"})
#     finally:
#         os.remove(tmp_path)

#     extracted = extract_quick_fields(text)

#     raw_json = {
#         "personal_info": {
#             "name": extracted["names"][0] if extracted.get("names") else "",
#             "emails": extracted.get("emails", []),
#             "phones": extracted.get("phones", [])
#         },
#         "locations": extracted.get("locations", []),
#         "education": extracted.get("education", []),
#         "experience": extracted.get("experience", []),
#         "skills": extracted.get("skills", []),
#         "projects": extracted.get("projects", []),
#         "achievements": extracted.get("achievements", []),
#         "extracurricular": extracted.get("extracurricular", [])
#     }

#     # Use Gemini to clean JSON and extract location properly
#     prompt = f"""
# You are given a resume JSON. Return the JSON cleaned and normalized.
# Split any location information into exact parts: 'address', 'city', 'state', 'country'.
# Keep all other fields as they are. Return only valid JSON, no markdown.
# Input JSON:
# {json.dumps(raw_json, indent=2)}
# """
#     async with httpx.AsyncClient(timeout=60) as client:
#         url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
#         payload = {"contents":[{"parts":[{"text": prompt}]}],"generationConfig":{"temperature":0,"maxOutputTokens":2000}}
#         resp = await client.post(url, json=payload)
#         resp.raise_for_status()
#         data = resp.json()

#     candidates = data.get("candidates", [])
#     content = ""
#     if candidates:
#         parts = candidates[0].get("content", {}).get("parts", [])
#         content = parts[0].get("text","") if parts else ""
#         content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()

#     try:
#         structured = json.loads(content)
#     except:
#         structured = raw_json  # fallback

#     # Ensure arrays are present
#     for key in ["education","experience","skills","projects","achievements","extracurricular","urls"]:
#         if key not in structured or structured[key] is None:
#             structured[key] = []

#     # Ensure location fields exist
#     for loc_field in ["address","city","state","country"]:
#         if loc_field not in structured:
#             structured[loc_field] = ""

#     return {"status":"ok", "resume_parsed": structured, "scanned": ext in ["png","jpg","jpeg"]}



@app.post("/login")
async def login_user(email: str = Form(...), password: str = Form(...)):
    user = USERS_DB.get(email)
    if not user or user["password"] != password:
        return JSONResponse(status_code=401, content={"status":"error","message":"Invalid credentials"})
    return {"status":"ok", "message":"Login successful","user":{"name":user["name"],"role":user["role"],"email":user["email"]}}