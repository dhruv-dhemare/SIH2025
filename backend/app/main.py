# fully working for pdf , docx
# import os
# import re
# import asyncio
# import pdfplumber
# from docx import Document
# from paddleocr import PaddleOCR
# from fastapi import FastAPI, File, UploadFile
# from fastapi.responses import JSONResponse
# import tempfile
# import httpx
# import json
# import spacy
# from dotenv import load_dotenv
# from PIL import Image, ImageEnhance, ImageFilter

# # ---------------------------
# # Load .env for API Key
# # ---------------------------
# load_dotenv()
# api_key = os.environ.get("GOOGLE_API_KEY")
# print(api_key)
# if not api_key:
#     raise ValueError("Set your GOOGLE_API_KEY in .env")

# # ---------------------------
# # App + Models Init
# # ---------------------------
# app = FastAPI()
# ocr = PaddleOCR(use_angle_cls=True, lang='en')
# nlp = spacy.load("en_core_web_sm")

# # ---------------------------
# # Helpers
# # ---------------------------
# def get_file_extension(filename: str):
#     return filename.split(".")[-1].lower()

# def normalize_list(values):
#     seen, cleaned = set(), []
#     for v in values:
#         val = v.strip()
#         if not val or val.lower() in seen:
#             continue
#         seen.add(val.lower())
#         cleaned.append(val)
#     return cleaned

# # ---------------------------
# # Extract Text
# # ---------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     text = []
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#     return "\n".join(text)

# def extract_text_from_docx(file_path: str) -> str:
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])



# def _preprocess_image_for_ocr(path: str) -> str:
#     """
#     Make a higher-contrast, larger temporary copy for OCR.
#     Returns path to preprocessed image (caller should remove it).
#     """
#     img = Image.open(path).convert("RGB")

#     # Resize up if small — OCR likes larger text
#     target_w = 1400
#     if img.width < target_w:
#         ratio = target_w / img.width
#         img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)

#     # Convert to gray + increase contrast + slight sharpen/smooth
#     gray = img.convert("L")
#     gray = ImageEnhance.Contrast(gray).enhance(1.6)
#     gray = gray.filter(ImageFilter.SMOOTH)

#     tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
#     tmp_path = tmp.name
#     tmp.close()
#     gray.save(tmp_path)
#     return tmp_path

# def extract_text_from_image(file_path: str) -> str:
#     # quick existence check
#     if not os.path.exists(file_path):
#         return f"ERROR: File not found at {file_path}"

#     # Preprocess and run OCR on preprocessed image first
#     pre_path = None
#     try:
#         pre_path = _preprocess_image_for_ocr(file_path)
#         results = ocr.ocr(pre_path)
#     except Exception as e:
#         # fallback to running on original image
#         results = None

#     # If preprocess gave nothing, try original
#     if not results or (isinstance(results, list) and len(results) == 0):
#         try:
#             results = ocr.ocr(file_path)
#         except Exception:
#             results = None

#     # Robustly extract text lines from many PaddleOCR return shapes
#     lines = []
#     if results:
#         # result is usually a list of lists: [ [ [box, (text, score), ...], ... ], ... ]
#         try:
#             for block in results:
#                 # block may be a list of items
#                 if isinstance(block, (list, tuple)):
#                     for item in block:
#                         # item usually like [box, (text, score)] or [box, text]
#                         try:
#                             candidate = item[1]
#                             # candidate may be tuple (text, prob) or text string
#                             if isinstance(candidate, (list, tuple)):
#                                 text = candidate[0]
#                             else:
#                                 text = candidate
#                             if text and isinstance(text, str) and text.strip():
#                                 lines.append(text.strip())
#                         except Exception:
#                             # fallback: try stringifying item
#                             try:
#                                 s = str(item)
#                                 if s.strip():
#                                     lines.append(s.strip())
#                             except:
#                                 continue
#                 else:
#                     # unexpected block shape, stringify
#                     s = str(block)
#                     if s.strip():
#                         lines.append(s.strip())
#         except Exception:
#             # last-resort: stringify whole results
#             lines = [str(results)]

#     # clean up temp preprocessed file
#     try:
#         if pre_path and os.path.exists(pre_path):
#             os.remove(pre_path)
#     except:
#         pass

#     if not lines:
#         return "ERROR: No text detected in image (PaddleOCR returned empty)."

#     return "\n".join(lines)

#     print("OCR text:\n", extract_text_from_image("resume.png"))




# # ---------------------------
# # Regex + spaCy Preprocessing
# # ---------------------------
# def extract_quick_fields(text: str):
#     # Emails and phones
#     emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
#     phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))

#     # spaCy NER
#     doc = nlp(text)
#     names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
#     orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
#     locations = normalize_list([ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]])

#     # Education (look for degrees, colleges, years)
#     edu_matches = re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I)
#     education_list = normalize_list(edu_matches)

#     # Skills (common keywords and tech stack pattern)
#     skills_matches = re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I)
#     skills_list = normalize_list(skills_matches)

#     # Projects (detect project-like titles and bullet points)
#     project_matches = re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text)
#     projects_list = normalize_list(project_matches)

#     # Experience (company + role + years)
#     experience_matches = re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text)
#     experience_list = normalize_list(["{} {} {}".format(*exp).strip() for exp in experience_matches if any(exp)])

#     # Achievements / Extracurricular
#     achievement_matches = re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text)
#     achievements_list = normalize_list(achievement_matches)

#     extracurricular_matches = re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text)
#     extracurricular_list = normalize_list(extracurricular_matches)

#     return {
#         "emails": emails,
#         "phones": phones,
#         "names": names,
#         "orgs": orgs,
#         "locations": locations,
#         "education": education_list,
#         "skills": skills_list,
#         "projects": projects_list,
#         "experience": experience_list,
#         "achievements": achievements_list,
#         "extracurricular": extracurricular_list
#     }


# # ---------------------------
# # Gemini Cleaning & Refinement (not parsing)
# # ---------------------------
# async def clean_with_gemini(raw_json: dict):
#     url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"

#     prompt = f"""
# You are a resume cleaning assistant.  
# The input is already parsed into JSON using regex, spaCy, and OCR.  
# Your task is ONLY to:
# - Remove irrelevant or duplicate values
# - Normalize names, degrees, companies, and skills
# - Ensure each section has relevant info only
# - Keep everything structured under correct keys

# ⚠️ IMPORTANT: Return ONLY valid JSON without markdown, without extra text, without ```json fences.

# Here is the JSON to clean:
# {json.dumps(raw_json, indent=2)}
# """

#     payload = {
#         "contents": [
#             {
#                 "parts": [
#                     {"text": prompt}
#                 ]
#             }
#         ],
#         "generationConfig": {
#             "temperature": 0,
#             "maxOutputTokens": 2000,
#         }
#     }

#     async with httpx.AsyncClient(timeout=60) as client:
#         response = await client.post(url, json=payload)
#         response.raise_for_status()
#         result = response.json()

#     candidates = result.get("candidates", [])
#     if not candidates:
#         return {"error": "No candidates returned from Gemini"}

#     # Gemini v1beta → candidates[0].content.parts[0].text
#     parts = candidates[0].get("content", {}).get("parts", [])
#     content = parts[0].get("text", "") if parts else ""

#     # Strip unwanted ```json fences
#     content = content.strip()
#     if content.startswith("```"):
#         content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
#         content = content.rstrip("```").strip()

#     try:
#         return json.loads(content)
#     except Exception:
#         return {"error": "Failed to parse cleaned JSON", "raw": content}

# # ---------------------------
# # API Endpoint
# # ---------------------------

# @app.post("/ocr-debug")
# async def ocr_debug(file: UploadFile = File(...)):
#     # write temp file
#     ext = file.filename.split(".")[-1].lower()
#     with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#         tmp.write(await file.read())
#         tmp_path = tmp.name

#     try:
#         # run OCR (no preprocessing) to inspect raw structure
#         raw = ocr.ocr(tmp_path)
#         # Return a safe string representation (some objects may not be JSON serializable)
#         os.remove(tmp_path)
#         return {"status": "ok", "ocr_raw": str(raw)}
#     except Exception as e:
#         try:
#             os.remove(tmp_path)
#         except:
#             pass
#         return {"status": "error", "message": str(e)}

# @app.post("/parse-resume")
# async def parse_resume_file(file: UploadFile = File(...)):
#     try:
#         ext = get_file_extension(file.filename)

#         with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
#             tmp.write(await file.read())
#             tmp_path = tmp.name

#         if ext == "pdf":
#             text = extract_text_from_pdf(tmp_path)
#         elif ext == "docx":
#             text = extract_text_from_docx(tmp_path)
#         elif ext in ["png", "jpg", "jpeg"]:
#             text = extract_text_from_image(tmp_path)
#         else:
#             os.remove(tmp_path)
#             return JSONResponse(
#                 status_code=400,
#                 content={"status": "error", "message": f"Unsupported file type: {ext}"}
#             )

#         os.remove(tmp_path)

#         # Step 1: Extract raw fields via regex + spaCy
#         extracted_data = extract_quick_fields(text)

#         raw_json = {
#             "personal_info": {
#                 "name": extracted_data["names"][0] if extracted_data["names"] else None,
#                 "emails": extracted_data["emails"],
#                 "phones": extracted_data["phones"],
#                 "locations": extracted_data["locations"],
#                 "organizations": extracted_data["orgs"]
#             },
#             "education": extracted_data["education"],
#             "experience": extracted_data["experience"],
#             "skills": extracted_data["skills"],
#             "projects": extracted_data["projects"],
#             "achievements": extracted_data["achievements"],
#             "extracurricular": extracted_data["extracurricular"]
#         }

#         # Step 2: Pass JSON to Gemini for structuring/cleaning
#         structured_data = await clean_with_gemini(raw_json)


#         return {
#             "status": "ok",
#             "resume_parsed": structured_data,
#             "scanned": ext in ["png", "jpg", "jpeg"]
#         }

#     except Exception as e:
#         return JSONResponse(
#             status_code=500,
#             content={"status": "error", "message": str(e)}
#         )


import os
import re
import asyncio
import pdfplumber
from docx import Document
from paddleocr import PaddleOCR
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
import tempfile
import httpx
import json
import spacy
from dotenv import load_dotenv
from PIL import Image, ImageEnhance, ImageFilter

# ---------------------------
# Load .env for API Key
# ---------------------------
load_dotenv()
api_key = os.environ.get("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("Set your GOOGLE_API_KEY in .env")

# ---------------------------
# App + Models Init
# ---------------------------
app = FastAPI()
ocr = PaddleOCR(use_angle_cls=True, lang='en')  # robust for rotated text
nlp = spacy.load("en_core_web_sm")

# ---------------------------
# Helpers
# ---------------------------
def get_file_extension(filename: str):
    return filename.split(".")[-1].lower()

def normalize_list(values):
    seen, cleaned = set(), []
    for v in values:
        val = v.strip()
        if not val or val.lower() in seen:
            continue
        seen.add(val.lower())
        cleaned.append(val)
    return cleaned

# ---------------------------
# Extract Text
# ---------------------------
def extract_text_from_pdf(file_path: str) -> str:
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# ---------------------------
# Image OCR Preprocessing
# ---------------------------
def _preprocess_image_for_ocr(path: str) -> str:
    img = Image.open(path).convert("RGB")
    target_w = 1400
    if img.width < target_w:
        ratio = target_w / img.width
        img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.6)
    gray = gray.filter(ImageFilter.SMOOTH)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = tmp.name
    tmp.close()
    gray.save(tmp_path)
    return tmp_path

def extract_text_from_image(file_path: str) -> str:
    if not os.path.exists(file_path):
        return f"ERROR: File not found at {file_path}"

    pre_path = None
    results = None
    try:
        pre_path = _preprocess_image_for_ocr(file_path)
        results = ocr.ocr(pre_path)
    except:
        results = None

    if not results or len(results) == 0:
        try:
            results = ocr.ocr(file_path)
        except:
            results = None

    lines = []
    if results:
        try:
            for block in results:
                if isinstance(block, (list, tuple)):
                    for item in block:
                        candidate = item[1]
                        if isinstance(candidate, (list, tuple)):
                            text = candidate[0]
                        else:
                            text = candidate
                        if text and isinstance(text, str) and text.strip():
                            lines.append(text.strip())
                else:
                    s = str(block)
                    if s.strip():
                        lines.append(s.strip())
        except:
            lines = [str(results)]

    if pre_path and os.path.exists(pre_path):
        os.remove(pre_path)

    if not lines:
        return "ERROR: No text detected in image (PaddleOCR returned empty)."
    return "\n".join(lines)

# ---------------------------
# Regex + spaCy Preprocessing
# ---------------------------
def extract_quick_fields(text: str):
    emails = normalize_list(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text))
    phones = normalize_list(re.findall(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text))
    doc = nlp(text)
    names = normalize_list([ent.text for ent in doc.ents if ent.label_ == "PERSON"])
    orgs = normalize_list([ent.text for ent in doc.ents if ent.label_ == "ORG"])
    locations = normalize_list([ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]])
    edu_matches = re.findall(r"(B\.?Tech|B\.?E|M\.?Tech|M\.?E|Bachelor|Master|Ph\.?D)[^•\n]*", text, re.I)
    education_list = normalize_list(edu_matches)
    skills_matches = re.findall(r"(Python|C\+\+|JavaScript|React|Node\.js|HTML|CSS|Tailwind|MongoDB|Firebase|Docker|GitHub|Flutter|Dart|SQL|MySQL|Vercel|Figma)", text, re.I)
    skills_list = normalize_list(skills_matches)
    project_matches = re.findall(r"(?:•\s)?([A-Z][A-Za-z0-9\s\-:]+(?:—[^•\n]+)?)", text)
    projects_list = normalize_list(project_matches)
    experience_matches = re.findall(r"(?:(?:at|with)\s)?([A-Z][A-Za-z0-9\s\-&]+)[\s,–]*(?:as|role)?\s*([A-Za-z\s]+)?[\s,–]*(\d{4}-\d{4}|\d{4})?", text)
    experience_list = normalize_list(["{} {} {}".format(*exp).strip() for exp in experience_matches if any(exp)])
    achievement_matches = re.findall(r"(?:•\s)?(Ranked|Awarded|Secured|Participated|Reached|Winner)[^\n]+", text)
    achievements_list = normalize_list(achievement_matches)
    extracurricular_matches = re.findall(r"(?:•\s)?(Volunteer|Coordinator|Member|Lead)[^\n]+", text)
    extracurricular_list = normalize_list(extracurricular_matches)

    return {
        "emails": emails,
        "phones": phones,
        "names": names,
        "orgs": orgs,
        "locations": locations,
        "education": education_list,
        "skills": skills_list,
        "projects": projects_list,
        "experience": experience_list,
        "achievements": achievements_list,
        "extracurricular": extracurricular_list
    }

# ---------------------------
# Gemini Cleaning (optional, not parsing)
# ---------------------------
async def clean_with_gemini(raw_json: dict):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    prompt = f"""
You are a resume cleaning assistant. Input is already parsed JSON.
Remove duplicates, irrelevant values, normalize names/degrees/skills.
Return only valid JSON, no markdown or ```json fences.
JSON:
{json.dumps(raw_json, indent=2)}
"""
    payload = {"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature":0, "maxOutputTokens":2000}}
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(url, json=payload)
        response.raise_for_status()
        result = response.json()
    candidates = result.get("candidates", [])
    if not candidates:
        return {"error": "No candidates returned from Gemini"}
    parts = candidates[0].get("content", {}).get("parts", [])
    content = parts[0].get("text", "") if parts else ""
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```[a-zA-Z]*\n?", "", content).rstrip("```").strip()
    try:
        return json.loads(content)
    except:
        return {"error": "Failed to parse cleaned JSON", "raw": content}

# ---------------------------
# API Endpoints
# ---------------------------
@app.post("/ocr-debug")
async def ocr_debug(file: UploadFile = File(...)):
    ext = get_file_extension(file.filename)
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        raw = ocr.ocr(tmp_path)
        os.remove(tmp_path)
        return {"status": "ok", "ocr_raw": str(raw)}
    except Exception as e:
        try: os.remove(tmp_path)
        except: pass
        return {"status": "error", "message": str(e)}

@app.post("/parse-resume")
async def parse_resume_file(file: UploadFile = File(...)):
    try:
        ext = get_file_extension(file.filename)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        if ext == "pdf":
            text = extract_text_from_pdf(tmp_path)
        elif ext == "docx":
            text = extract_text_from_docx(tmp_path)
        elif ext in ["png", "jpg", "jpeg"]:
            text = extract_text_from_image(tmp_path)
        else:
            os.remove(tmp_path)
            return JSONResponse(status_code=400, content={"status": "error", "message": f"Unsupported file type: {ext}"})

        os.remove(tmp_path)
        extracted_data = extract_quick_fields(text)

        raw_json = {
            "personal_info": {
                "name": extracted_data["names"][0] if extracted_data["names"] else None,
                "emails": extracted_data["emails"],
                "phones": extracted_data["phones"],
                "locations": extracted_data["locations"],
                "organizations": extracted_data["orgs"]
            },
            "education": extracted_data["education"],
            "experience": extracted_data["experience"],
            "skills": extracted_data["skills"],
            "projects": extracted_data["projects"],
            "achievements": extracted_data["achievements"],
            "extracurricular": extracted_data["extracurricular"]
        }

        structured_data = await clean_with_gemini(raw_json)
        return {"status": "ok", "resume_parsed": structured_data, "scanned": ext in ["png", "jpg", "jpeg"]}
    except Exception as e:
        return JSONResponse(status_code=500, content={"status": "error", "message": str(e)})
